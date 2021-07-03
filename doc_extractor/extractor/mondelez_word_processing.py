import pandas as pd
import re
from langid import classify
import numpy as np
from langid import classify
from sklearn.neural_network import MLPClassifier
import time
from laserembeddings import Laser
import warnings
warnings.filterwarnings("ignore")
import joblib
from docx import Document
import io
import smbclient
import tempfile

from environment import MODE

if MODE == 'local':
    from .local_constants import *
else:
    from .dev_constants import *

laser = Laser(path_to_bpe_codes, path_to_bpe_vocab, path_to_encoder)

nutrition_keywors = ['Energy','Calories','Protein','Carbohydrate','Sugars','Total Fat','Saturated Fat','Fibre',
                    'Sodium','Salt']

nutri_keys_para = ['total fat','saturated fat','transfat','cholesterol','sodium','carbohydrate',
                  'dietary fiber','total sugars','includes added sugars','protein','servings per package',
                  'serving size','calories','servings per container']

# mondlelz_word_model_location = r"/Users/sakthivel/Documents/SGK/Mondelez-Word/Dataset/mondelez_word_model.sav"

classifier = joblib.load(mondelez_word_model_location)

def get_input(input_file,input_docx_location):
    if input_file.startswith('\\'):
        print('connecting to SMB share')
        try:
            with smbclient.open_file(r"{}".format(input_file), mode='rb', username=smb_username,
                                     password=smb_password) as f:
                with open(input_docx_location, 'wb') as pdf:
                    pdf.write(f.read())
                print('file found')
        except:
            smbclient.reset_connection_cache()
            with smbclient.open_file(r"{}".format(input_file), mode='rb', username=smb_username,
                                     password=smb_password) as f:
                with open(input_docx_location, 'wb') as pdf:
                    pdf.write(f.read())
                print('file found')
        finally:
            smbclient.reset_connection_cache()
        return input_docx_location
    else:
        return document_location + input_file

def remove_duplicates(dictionary):
    final_cleaned_dict = {}
    for category , value_list in dictionary.items():
        final_cleaned_dict[category] = sorted(list({frozenset(list_element.keys()) : list_element for list_element in value_list}.values()),key=lambda d: list(d.keys()))
    return final_cleaned_dict

def preprocess(file):
    document = Document(file)

    table_content_list = []
    tables = document.tables
    for table in tables:
        for row in table.rows:
            row_list = []
            for cell in row.cells:
                row_list.append(str(cell.text))
            table_content_list.append(row_list)

    table_content_list_new = []
    for i in range(0, len(table_content_list)):
        row_number = [table_content_list[i], i + 1]
        table_content_list_new.append(row_number)

    header_list = []
    for cnt in range(0, len(table_content_list_new)):
        sr = pd.Series(table_content_list_new[cnt][0])
        result = sr.first_valid_index()
        if sr[result].strip() != '':
            values = [sr[result].replace('\xa0', ''), table_content_list_new[cnt][1]]
            header_list.append(values)
        else:
            #         print()
            content = str(('').join(table_content_list_new[cnt][0])).strip()
            if content[-4:].lower() == 'kcal':
                values = ['calories', table_content_list_new[cnt][1]]
                #             print(values)
                header_list.append(values)

    prediction_list = []
    ingre_dic = {}
    nutri_para_list = []
    for item in header_list:
        #     print(item[0])
        classified_output = classifier.predict(laser.embed_sentences(item[0], lang='en'))
        probability1 = classifier.predict_proba(laser.embed_sentences(item[0], lang='en'))
        probability1.sort()
        prob1 = probability1[0][-1]
        if prob1 > 0.75:
            #         if (prob1 > 0.75) or ((prob1 / 2) > probability1[0][-2]):
            classified_output = classified_output[0]
        else:
            classified_output = 'None'
        if classified_output not in ['None', 'INGREDIENTS_DECLARATION', 'NUTRITION_INSTRUCTIONS',
                                     'NUTRITION_PARAGRAPH']:
            contents = [item[0], classified_output.strip(), item[1]]
            prediction_list.append(contents)

        elif classified_output in ['INGREDIENTS_DECLARATION', 'NUTRITION_INSTRUCTIONS']:

            lang = classify(item[0])[0]
            if classified_output in ingre_dic:
                ingre_dic[classified_output].append({lang: item[0]})
            else:
                ingre_dic[classified_output] = [{lang: item[0]}]

        elif classified_output in ['NUTRITION_PARAGRAPH']:
            nutri_para_list.append(item[0])

    ingre_new_dic = {}
    for k1, v1 in ingre_dic.items():
        for cnts in v1:
            for lang, txt in cnts.items():
                split_txt = (txt.split('\n\n'))
                #             print('*'*10)
                for txts in split_txt:
                    if txts.strip() not in ['', '\n']:
                        classified_output = classifier.predict(laser.embed_sentences(txts, lang='en'))
                        probability1 = classifier.predict_proba(laser.embed_sentences(txts, lang='en'))
                        probability1.sort()
                        prob1 = probability1[0][-1]
                        if prob1 > 0.75:
                            classified_output = classified_output[0]
                        else:
                            classified_output = 'None'

                        if classified_output != 'None':
                            if classified_output in ingre_new_dic:
                                ingre_new_dic[classified_output].append({lang: txts})
                            else:
                                ingre_new_dic[classified_output] = [{lang: txts}]

                        elif classified_output == 'None':
                            if k1 in ingre_new_dic:
                                ingre_new_dic[k1].append({lang: txts})
                            else:
                                ingre_new_dic[k1] = [{lang: txts}]

    final_content_list = []
    nutri_content_list = []
    gen_list = []
    nutri_list = []

    for l in range(0, len(table_content_list_new)):
        for m in range(0, len(prediction_list)):
            if prediction_list[m][1] not in nutrition_keywors:
                if prediction_list[m][2] == table_content_list_new[l][1]:
                    cont = [table_content_list_new[l][0], prediction_list[m][1]]
                    final_content_list.append(cont)
                    gen_list.append(prediction_list[m][1])
            elif prediction_list[m][1] in nutrition_keywors:
                if prediction_list[m][2] == table_content_list_new[l][1]:
                    cont1 = [table_content_list_new[l][0], prediction_list[m][1]]
                    nutri_content_list.append(cont1)
                    nutri_list.append(prediction_list[m][1])

    gen_cate_dic = {}
    for p in range(0,len(final_content_list)):
        if len(final_content_list[p][0]) > 1:
            for d in range(1,len(final_content_list[p][0])):
                if final_content_list[p][1] in ['Legal_Designation','LOCATION_OF_ORIGIN','VARIANT','Servings Per Package','Serving Size','BRAND_NAME','Serving Per Container','NET_CONTENT_STATEMENT']:
                    split_cnt = final_content_list[p][0][d].split('\n')
                    for texts in split_cnt:
                        if texts.strip()!= '':
                            lang = classify(texts)[0]
                            if final_content_list[p][1] in gen_cate_dic:
                                gen_cate_dic[final_content_list[p][1]].append({lang:texts.strip()})
                            else:
                                gen_cate_dic[final_content_list[p][1]]= [{lang:texts.strip()}]
                else:
                    lang = classify(final_content_list[p][0][d])[0]
                    if final_content_list[p][1] in gen_cate_dic:
                        gen_cate_dic[final_content_list[p][1]].append({lang:final_content_list[p][0][d]})
                    else:
                        gen_cate_dic[final_content_list[p][1]]= [{lang:final_content_list[p][0][d]}]

        elif len(final_content_list[p][0]) ==1:
            lang = classify(final_content_list[p][0][0])[0]
            if final_content_list[p][1] in gen_cate_dic:
                gen_cate_dic[final_content_list[p][1]].append({lang:final_content_list[p][0][0]})
            else:
                gen_cate_dic[final_content_list[p][1]]= [{lang:final_content_list[p][0][0]}]

    for i in range(0, len(table_content_list)):
        if 'brand' in table_content_list[i][0].lower():
            brand_indx = i
            if len(table_content_list[brand_indx]) == 2:
                if table_content_list[brand_indx + 1][0].strip() == '':
                    if table_content_list[brand_indx + 1][1].strip() != '':
                        lang = classify(table_content_list[brand_indx + 1][1])[0]
                        if "BRAND_NAME" in gen_cate_dic:
                            gen_cate_dic["BRAND_NAME"].append({lang: table_content_list[brand_indx + 1][1]})
                        else:
                            gen_cate_dic["BRAND_NAME"] = [{lang: table_content_list[brand_indx + 1][1]}]

    for k3, v3 in ingre_new_dic.items():                # Allergen statement issue fixed
        if k3 == 'ALLERGEN_STATEMENT':
            for cnt_list in v3:
                for lang, texts in cnt_list.items():
                    if k3 in gen_cate_dic:
                        gen_cate_dic[k3].append({lang: texts})
                    else:
                        gen_cate_dic[k3] = [{lang: texts}]

    ingre_new_dic.pop('ALLERGEN_STATEMENT', None)

    return table_content_list, gen_cate_dic, nutri_para_list, nutri_content_list, nutri_list, ingre_new_dic


def paragraph(file):
    document = Document(file)

    paragraph_txt = []
    all_paras = document.paragraphs
    for para in all_paras:
        paragraph_txt.append(para.text.strip())

    paragraph_txt = [txt.strip() for txt in paragraph_txt if txt]

    para_dic = {}
    for txt in paragraph_txt:
        classified_output = classifier.predict(laser.embed_sentences(txt, lang='en'))
        probability1 = classifier.predict_proba(laser.embed_sentences(txt, lang='en'))
        probability1.sort()
        prob1 = probability1[0][-1]
        if prob1 > 0.75:
            classified_output = classified_output[0]
        else:
            classified_output = 'None'

        if classified_output == "NUTRITION_TABLE_CONTENT":

            lang = classify(txt)[0]
            if "NUTRITION_TABLE_CONTENT" in para_dic:
                para_dic["NUTRITION_TABLE_CONTENT"].append({lang: txt})
            else:
                para_dic["NUTRITION_TABLE_CONTENT"] = [{lang: txt}]

    return para_dic


def nutrition(dictionary):
    nutri_dic = {}
    for keys, value in dictionary.items():

        classified_output = classifier.predict(laser.embed_sentences(keys.replace('\xa0', ''), lang='en'))
        probability1 = classifier.predict_proba(laser.embed_sentences(keys.replace('\xa0', ''), lang='en'))
        probability1.sort()
        prob1 = probability1[0][-1]
        if prob1 > 0.75:
            classified_output = classified_output[0]
        else:
            classified_output = 'None'

        if classified_output != 'None':
            if classified_output in nutri_dic:

                nutri_dic[classified_output].append(value)
            else:
                nutri_dic[classified_output] = value

        else:

            if keys in nutri_dic:

                nutri_dic[keys].append(value)
            else:
                nutri_dic[keys] = value

    return nutri_dic


def serving_dictionary(dic, gen_cate_dic):
    for k, v in dic.items():
        classified_output = classifier.predict(laser.embed_sentences(k.replace('\xa0', ''), lang='en'))
        probability1 = classifier.predict_proba(laser.embed_sentences(k.replace('\xa0', ''), lang='en'))
        probability1.sort()
        prob1 = probability1[0][-1]
        if prob1 > 0.75:
            classified_output = classified_output[0]
        else:
            classified_output = 'None'

        if classified_output != 'None':
            if classified_output in gen_cate_dic:

                gen_cate_dic[classified_output].append(v)
            else:
                gen_cate_dic[classified_output] = v
        else:
            if k in gen_cate_dic:

                gen_cate_dic[k].append(values)
            else:
                gen_cate_dic[k] = values

    return gen_cate_dic


def nutrition_para_format(nutri_para_list, gen_cate_dic, table_content_list):
    content_list = []
    for i in nutri_para_list:
        list_of_cnt = i.split(',')
        content_list.append(list_of_cnt)

    nutri_dic = []
    for k1 in range(0, len(content_list)):

        nutrit_para_dic = {}
        #         nutri_dic_new={}
        d_regex_list = []

        for i in content_list[k1]:
            rege = ('').join(re.findall(r'[<.a-zA-Z0-9-(%)]', str(i)))
            d_regex_list.append(rege)

        d_regex_new_list = []
        for j in d_regex_list:
            reg1 = re.findall(r'(<?\s?(\d?\d)(\.\d+)?\s?(mg|kj|g|%|mcg))', str(j))
            d_regex_new_list.append(reg1)
            #             print(d_regex_new_list)

        for l5 in range(0, len(d_regex_new_list)):
            for m5 in range(0, len(d_regex_new_list[l5])):
                for keys in nutri_keys_para:
                    if keys in content_list[k1][l5].lower():
                        if '%' in d_regex_new_list[l5][m5][0]:
                            if keys in nutrit_para_dic:
                                nutrit_para_dic[keys].append({'PDV': {'en': d_regex_new_list[l5][m5][0]}})
                            else:
                                nutrit_para_dic[keys] = [{'PDV': {'en': d_regex_new_list[l5][m5][0]}}]
                        else:
                            if keys in nutrit_para_dic:
                                nutrit_para_dic[keys].append({'Values': {'en': d_regex_new_list[l5][m5][0]}})
                            else:
                                nutrit_para_dic[keys] = [{'Values': {'en': d_regex_new_list[l5][m5][0]}}]

        nutri_dic_new = nutrition(nutrit_para_dic)

        nutri_dic.append(nutri_dic_new)

    serving_list = []
    serving_key = ["servings per container", "amount per serving"]
    for i in range(0, len(table_content_list)):
        for k in serving_key:
            if k in table_content_list[i][0].lower():
                serving_list.append(table_content_list[i][0].replace('\xa0', ''))

    final_serving_list = []
    for j in range(0, len(serving_list)):
        split_list = serving_list[j].split(',')
        final_serving_list.append(split_list)

    final_serving_list = sum(final_serving_list, [])

    para_regex_list = []
    para_regex_new_list = []
    for i in final_serving_list:
        #     reg = re.findall(r'(<?\s+(\s?\d)(\.\d+)?\s?(mg|kj|g|%|mcg))|(<?\s+(\.\d+)?\s?(mg|kj|g|%|mcg))',str(i))
        para_reg = ('').join(re.findall(r'[<.a-zA-Z0-9-(%)]', str(i)))
        para_regex_list.append(para_reg)

        # regex_list

    for j in para_regex_list:
        para_reg1 = re.findall(r'(<?\s?(\d?\d)(\.\d+)?\s?(mg|kj|g|%|mcg))|(\d?\d?\d)', str(j))
        para_regex_new_list.append(para_reg1)

    para_tuple_list = []
    for m in para_regex_new_list:
        para_tuple_list.append([(tuple(int(x) if x.isdigit() else x for x in _ if x)) for _ in m])

    serving_dic_items = {}
    serving_dic = {}
    for l in range(0, len(para_tuple_list)):
        for m in range(0, len(para_tuple_list[l])):
            for k in nutri_keys_para:
                if k in final_serving_list[l].replace('\x95', '').replace('\xa0', '').lower():
                    if '%' in str(para_tuple_list[l][m][0]):
                        if k in serving_dic_items:
                            # serving_dic_items[k].append({'PDV': {'en': str(para_tuple_list[l][m][0])}})
                            serving_dic_items[k].append({'en': str(para_tuple_list[l][m][0])})
                        else:
                            # serving_dic_items[k] = [{'PDV': {'en': str(para_tuple_list[l][m][0])}}]
                            serving_dic_items[k] = [{'en': str(para_tuple_list[l][m][0])}]
                    else:
                        if k in serving_dic_items:
                            # serving_dic_items[k].append({'Value': {'en': str(para_tuple_list[l][m][0])}})
                            serving_dic_items[k].append({'en': str(para_tuple_list[l][m][0])})
                        else:
                            # serving_dic_items[k] = [{'Value': {'en': str(para_tuple_list[l][m][0])}}]
                            serving_dic_items[k] = [{'en': str(para_tuple_list[l][m][0])}]

    gen_cate_dic = serving_dictionary(serving_dic_items, gen_cate_dic)
    gen_cate_dic.pop('Calories',None)
    return gen_cate_dic, nutri_dic

def table_nutrition_format(nutri_content_list):
    serving_dic={}
    nutri_dic ={}
    for n in range(0,len(nutri_content_list)):
        if len(nutri_content_list[n][0])>1:
            for o in range(1,len(nutri_content_list[n][0])):
        #         for keys in nutri_list:
        #             if keys in nutri_content_list[n][1]:
                        if nutri_content_list[n][0][o].strip() not in ['']:
                            if '%' in nutri_content_list[n][0][o]:
                                if nutri_content_list[n][1] in nutri_dic:
                                    nutri_dic[nutri_content_list[n][1]].append({'PDV':{'en':nutri_content_list[n][0][o].strip()}})
                                else:
                                    nutri_dic[nutri_content_list[n][1]] = [{'PDV':{'en':nutri_content_list[n][0][o].strip()}}]
                            else:

                                if nutri_content_list[n][1] in nutri_dic:
                                    nutri_dic[nutri_content_list[n][1]].append({'Values':{'en':nutri_content_list[n][0][o].strip()}})
                                else:
                                    nutri_dic[nutri_content_list[n][1]] = [{'Values':{'en':nutri_content_list[n][0][o].strip()}}]
    return nutri_dic


def non_table_nutrition(non_table):
    nutri_keys = []
    nutri_values = []
    for j in range(0, len(non_table)):
        for k in range(0, len(non_table[j])):
            if 'energy' in non_table[j][k].lower():
                nutri_keys.append(non_table[j][k].split('\n'))
            else:
                if '\n' in non_table[j][k].lower():
                    nutri_values.append(non_table[j][k].split('\n'))

    new_nutri_dic = {}
    serving_dic = {}
    if len(nutri_keys[0]) == len(nutri_values[0]):
        for j1 in range(0, len(nutri_values)):
            for j2 in range(0, len(nutri_values[j1])):
                if nutri_values[j1][j2].strip() != '':
                    if nutri_keys[0][j2] in new_nutri_dic:
                        if '%' in nutri_values[j1][j2].strip():
                            new_nutri_dic[nutri_keys[0][j2]].append({'PDV': {'en': nutri_values[j1][j2].strip()}})
                        else:
                            if 'kcal' in nutri_values[j1][j2].strip().lower():
                                new_nutri_dic['calories'].append({'Values': {'en': nutri_values[j1][j2].strip()}})
                            else:
                                new_nutri_dic[nutri_keys[0][j2]].append(
                                    {'Values': {'en': nutri_values[j1][j2].strip()}})
                    else:
                        if '%' in nutri_values[j1][j2].strip():
                            new_nutri_dic[nutri_keys[0][j2]] = [{'PDV': {'en': nutri_values[j1][j2].strip()}}]
                        else:

                            if 'kcal' in nutri_values[j1][j2].strip().lower():
                                if 'calories' in new_nutri_dic:
                                    new_nutri_dic['calories'].append({'Values': {'en': nutri_values[j1][j2].strip()}})
                                else:
                                    new_nutri_dic['calories'] = [{'Values': {'en': nutri_values[j1][j2].strip()}}]
                            else:
                                new_nutri_dic[nutri_keys[0][j2]] = [{'Values': {'en': nutri_values[j1][j2].strip()}}]

    nutri_dic = nutrition(new_nutri_dic)

    return nutri_dic

def arabic_table_nutrition(table_content_list, gen_cate_dic):
    start_indx,end_indx = 0,0
    for i in range(0, len(table_content_list)):
        if 'nutrition facts' in table_content_list[i][0].lower():
            start_indx = i + 1
        elif "% daily value" in table_content_list[i][0].lower():
            end_indx = i

    new_list = []
    for j in range(start_indx, end_indx):
        new_list.append(table_content_list[j])

    dupli_list = []
    for i in new_list:
        dupli_list.append(('').join(i))

    regex_list_ = []
    for i in dupli_list:
        #     reg = re.findall(r'(<?\s+(\s?\d)(\.\d+)?\s?(mg|kj|g|%|mcg))|(<?\s+(\.\d+)?\s?(mg|kj|g|%|mcg))',str(i))
        reg_ = ('').join(re.findall(r'[<.a-zA-Z0-9-(%)]', str(i).replace(',', '.')))
        regex_list_.append(reg_)
        # regex_list

    regex_new_list_ = []
    for j in regex_list_:
        #     reg1_ = re.findall(r'(<?\s?(\d?\d?\d)(\.\d+)?\s?(mg|kj|g|%|mcg))|(\d?\d?\d)',str(j).replace(',','.'))
        reg1_ = re.findall(r'(<?\s?(\d?\d?\d)(\.\d+)?(\,\d+)?\.?\s?(mg|kj|g|%|mcg))|(\d?\d?\d)',
                           str(j).replace(',', '.'))
        regex_new_list_.append(reg1_)

    tuple_list = []
    for m in regex_new_list_:
        tuple_list.append([(tuple(int(x) if x.isdigit() else x for x in _ if x)) for _ in m])

    regex_new_list_ = tuple_list

    nutri_arabic_dic = {}
    serving_dic = {}
    for l in range(0, len(regex_new_list_)):
        for m in range(0, len(regex_new_list_[l])):
            for k in nutri_keys_para:
                if k in dupli_list[l].lower().replace('\x95', '').replace('\xa0', ''):
                    if '%' in str(regex_new_list_[l][m][0]):
                        if k in nutri_arabic_dic:
                            nutri_arabic_dic[k].append({'PDV': {'en': str(regex_new_list_[l][m][0])}})
                        else:
                            nutri_arabic_dic[k] = [{'PDV': {'en': str(regex_new_list_[l][m][0])}}]
                    else:
                        if k in nutri_arabic_dic:
                            nutri_arabic_dic[k].append({'Values': {'en': str(regex_new_list_[l][m][0])}})
                        else:
                            nutri_arabic_dic[k] = [{'Values': {'en': str(regex_new_list_[l][m][0])}}]

    duplicate_dic = remove_duplicates(nutri_arabic_dic)

    nutri_dic = nutrition(duplicate_dic)

    serving_dic_items = {}
    serving_dic = {}
    serv_list = ['serving size', 'servings per container', 'servings per package']
    #     for k6,v6 in duplicate_dic.items():
    #         for ke in serv_list:
    #             if ke in k6.lower():
    #                 if ke in serving_dic_items:
    #                       serving_dic_items[ke].append(v6)
    #                 else:
    #                     serving_dic_items[ke] = v6
    for k, v in duplicate_dic.items():
        for txt in v:
            for k7, v7 in txt.items():
                for ke in serv_list:
                    if ke in k.lower():
                        if ke in serving_dic_items:
                            serving_dic_items[ke].append(v7)
                        else:
                            serving_dic_items[ke] = v7

    gen_cate_dic = serving_dictionary(serving_dic_items, gen_cate_dic)

    nutri_dic.pop('Serving Size', None)
    nutri_dic.pop('Serving Per Container', None)
    nutri_dic.pop('Servings Per Package', None)
    nutri_dic.pop('Servings Per Container', None)

    return nutri_dic, gen_cate_dic
#             nutri_dic.pop('servings per package', None)


def mondelez_word(file):
    t5 = time.time()

    temp_directory = tempfile.TemporaryDirectory(dir=document_location)
    input_docx_location = f'{temp_directory.name}/input_docx.docx'

    file = get_input(file,input_docx_location)

    [table_content_list, gen_cate_dic, nutri_para_list,
     nutri_content_list, nutri_list, ingre_new_dic] = preprocess(file)

    para_dic = paragraph(file)

    if nutri_para_list:
        gen_cate_dic, nutrition_dic = nutrition_para_format(nutri_para_list, gen_cate_dic, table_content_list)

    elif not nutri_list:
        non_table = []
        try:
            for i in range(0, len(table_content_list)):
                for j in range(0, len(table_content_list[i])):
                    if 'nutrition information' in table_content_list[i][j].lower():
                        if 'FOP icon' in table_content_list[i][j]:
                            break
                        list_content = [table_content_list[i], table_content_list[i + 1]]
                        non_table.append(list_content)
                    elif 'energy' in table_content_list[i][j].lower() and 'carbohydrate' in table_content_list[i][j].lower():
                        list_content = [table_content_list[i]]
                        non_table.append(list_content)
            non_table = sum(non_table, [])
        except:
            for i in range(0, len(table_content_list)):
                for j in range(0, len(table_content_list[i])):
                    if 'nutrition information' in table_content_list[i][j].lower():
                        if 'FOP icon' in table_content_list[i][j]:
                            break
                        list_content = [table_content_list[i]]
                        non_table.append(list_content)
                    elif 'energy' in table_content_list[i][j].lower() and 'carbohydrate' in table_content_list[i][j].lower():
                        list_content = [table_content_list[i]]
                        non_table.append(list_content)
            non_table = sum(non_table, [])

        if not non_table:
            nutrition_dic, gen_cate_dic = arabic_table_nutrition(table_content_list, gen_cate_dic)
        else:

            nutrition_dic = non_table_nutrition(non_table)


    else:
        nutrition_dic = table_nutrition_format(nutri_content_list)

    over_all_dic = {**gen_cate_dic, **ingre_new_dic, **para_dic}

    final_dic = {}
    if 'NUTRITION_FACTS' in final_dic:
        if isinstance(nutrition_dic, list):
            final_dic['NUTRITION_FACTS'].extend(nutrition_dic)
        else:
            final_dic['NUTRITION_FACTS'].append(nutrition_dic)
    else:
        if isinstance(nutrition_dic, list):
            final_dic['NUTRITION_FACTS'] = nutrition_dic
        else:
            final_dic['NUTRITION_FACTS'] = [nutrition_dic]
    #     final_dic['NUTRITION_FACTS'] = nutrition_dic
    t6 = time.time()
    print(f'Finished in {t6 - t5}seconds')
    return {**final_dic, **over_all_dic}
